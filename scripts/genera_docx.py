# -*- coding: utf-8 -*-
"""Generatore DOCX per i documenti dell'opera ITALIA NERA (CLAUDE.md §5).

Uso:  python3 scripts/genera_docx.py <corpo_body.py> <output.docx>

Il corpo del documento sta in un file `*_body.py` separato (mai heredoc: distrugge i
diacritici) che espone: TITOLO (str, opzionale), PARAGRAFI (list[str] di prosa, con
grassetto inline `**...**`) e, per le monografie, FONTI (list[str] di URL).

Vincoli applicati: A4, margini 2,54 cm, Times New Roman 12pt a livello di stile e di run
(w:rFonts con ascii/hAnsi/cs/eastAsia), giustificato, rientro prima riga 0,75 cm,
interlinea 1,15, spaziatura zero, zoom 100%. Prima della consegna esegue il controllo
«niente cifre arabe residue» (§5) e riapre il file per validarlo."""

import importlib.util
import re
import sys

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = "Times New Roman"


def _carica_body(path):
    spec = importlib.util.spec_from_file_location("body", path)
    mod = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(mod)
    return mod


def _imposta_rfonts(rpr, nome=FONT):
    rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts')
        rpr.insert(0, rf)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rf.set(qn(attr), nome)


def _aggiungi_prosa(par, testo):
    # segmenti di indice dispari (split su **) = grassetto
    for i, seg in enumerate(testo.split("**")):
        if not seg:
            continue
        run = par.add_run(seg)
        run.font.name = FONT
        run.font.size = Pt(12)
        run.bold = (i % 2 == 1)
        _imposta_rfonts(run._element.get_or_add_rPr())


def _paragrafo(doc, testo):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.first_line_indent = Cm(0.75)
    pf.line_spacing = 1.15
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    _aggiungi_prosa(p, testo)
    return p


def _zoom_100(doc):
    settings = doc.settings.element
    for z in settings.findall(qn('w:zoom')):
        settings.remove(z)
    z = OxmlElement('w:zoom')
    z.set(qn('w:percent'), "100")
    settings.insert(0, z)


def _controllo_cifre(paragrafi):
    """Restituisce la lista delle cifre arabe residue (deve essere vuota, §5)."""
    residui = []
    for testo in paragrafi:
        residui += re.findall(r'(?<![A-Za-z])\d+', testo)
    return residui


def genera(body_path, out_path, consenti_cifre=False):
    b = _carica_body(body_path)
    paragrafi = list(getattr(b, "PARAGRAFI", []))

    residui = _controllo_cifre(paragrafi)
    if residui and not consenti_cifre:
        raise SystemExit(
            "Cifre arabe residue nel testo dell'opera: %r. Scrivile per esteso "
            "(skill numerali-italiani) o dichiarale come citazioni diagnostiche "
            "(rilancia con consenti_cifre=True)." % residui)

    doc = Document()

    # Times New Roman 12pt a livello di STILE (Normal)
    normal = doc.styles['Normal']
    normal.font.name = FONT
    normal.font.size = Pt(12)
    _imposta_rfonts(normal.element.get_or_add_rPr())

    sec = doc.sections[0]
    sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
    sec.top_margin = sec.bottom_margin = Cm(2.54)
    sec.left_margin = sec.right_margin = Cm(2.54)

    for testo in paragrafi:
        _paragrafo(doc, testo)

    # monografia: sezione obbligatoria «Fonti e URL» (§1)
    fonti = getattr(b, "FONTI", None)
    if fonti:
        _paragrafo(doc, "**Fonti e URL**")
        for url in fonti:
            _paragrafo(doc, url)

    _zoom_100(doc)
    doc.save(out_path)

    Document(out_path)  # validazione: deve riaprirsi senza eccezioni
    return out_path


if __name__ == "__main__":
    if len(sys.argv) < 3:
        raise SystemExit(__doc__)
    print(genera(sys.argv[1], sys.argv[2]))
