# -*- coding: utf-8 -*-
"""Generatore DOCX per l'opera monumentale ITALIA NERA - variante a CIFRE ARABE.

Mantiene lo stile tipografico della sezione 5 del CLAUDE.md (A4, margini 2,54 cm,
Times New Roman 12 pt a livello di stile e di run, giustificato, rientro prima riga
0,75 cm, interlinea 1,15, zoom 100%) MA, per esplicita richiesta dell'autore su questo
documento, NON applica il guardiano 'niente cifre arabe': i numerali arabi sono ammessi.

Uso:  python3 gen_docx_cifre.py <corpo_body.py> <output.docx>
Il body espone: TITOLO (str), CONTENUTO (list di tuple ('h1'|'h2'|'p', testo)) e
FONTI (list[str]).  Il grassetto inline si segna con **...**.
"""
import importlib.util, re, sys
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = "Times New Roman"

def _carica(path):
    spec = importlib.util.spec_from_file_location("body", path)
    mod = importlib.util.module_from_spec(spec); spec.loader.exec_module(mod); return mod

def _rfonts(rpr, nome=FONT):
    rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rpr.insert(0, rf)
    for a in ('w:ascii','w:hAnsi','w:cs','w:eastAsia'): rf.set(qn(a), nome)

def _run(par, testo, bold=False, size=12):
    r = par.add_run(testo); r.bold = bold; r.font.name = FONT; r.font.size = Pt(size)
    _rfonts(r._element.get_or_add_rPr())
    return r

def _add_bold(par, testo, size=12):
    # gestisce **grassetto** inline
    for i, seg in enumerate(re.split(r'(\*\*.*?\*\*)', testo)):
        if not seg: continue
        if seg.startswith('**') and seg.endswith('**'):
            _run(par, seg[2:-2], bold=True, size=size)
        else:
            _run(par, seg, bold=False, size=size)

def main():
    body_path, out_path = sys.argv[1], sys.argv[2]
    b = _carica(body_path)
    doc = Document()
    st = doc.styles['Normal']; st.font.name = FONT; st.font.size = Pt(12)
    _rfonts(st.element.get_or_add_rPr())
    pf = st.paragraph_format; pf.line_spacing = 1.15; pf.space_before = Pt(0); pf.space_after = Pt(0)
    sec = doc.sections[0]
    sec.page_width = Cm(21); sec.page_height = Cm(29.7)
    sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Cm(2.54)

    titolo = getattr(b, 'TITOLO', None)
    if titolo:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0); p.paragraph_format.space_after = Pt(18)
        _run(p, titolo, bold=True, size=16)

    nwords = 0
    for tipo, testo in b.CONTENUTO:
        nwords += len(testo.split())
        if tipo == 'h1':
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_before = Pt(18); p.paragraph_format.space_after = Pt(6)
            _run(p, testo, bold=True, size=14)
        elif tipo == 'h2':
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(4)
            _run(p, testo, bold=True, size=12.5)
        else:
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Cm(0.75)
            _add_bold(p, testo)

    fonti = getattr(b, 'FONTI', None)
    if fonti:
        p = doc.add_paragraph(); p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_before = Pt(18); p.paragraph_format.space_after = Pt(6)
        _run(p, "Apparato delle fonti", bold=True, size=14)
        for u in fonti:
            pp = doc.add_paragraph(); pp.paragraph_format.first_line_indent = Cm(0)
            _run(pp, u, bold=False, size=10)

    # zoom 100%
    try:
        settings = doc.settings.element
        zoom = OxmlElement('w:zoom'); zoom.set(qn('w:percent'), '100'); settings.append(zoom)
    except Exception: pass

    doc.save(out_path)
    # validazione: riapertura
    Document(out_path)
    print(f"{out_path}  |  parole (corpo): {nwords}")

if __name__ == '__main__':
    main()
